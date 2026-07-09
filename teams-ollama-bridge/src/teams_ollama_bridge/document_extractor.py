"""Textextraktion aus Dokumenten."""

from __future__ import annotations

from pathlib import Path

from teams_ollama_bridge.exceptions import AttachmentExtractionError, EncryptedPdfError
from teams_ollama_bridge.logging_config import get_logger

logger = get_logger(__name__)


def _truncate_text(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."


def _read_text_file(path: Path, max_chars: int) -> str:
    raw = path.read_bytes()
    if raw.startswith(b"\xef\xbb\xbf"):
        text = raw.decode("utf-8-sig")
    else:
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("cp1252", errors="replace")
    return _truncate_text(text.strip(), max_chars)


def _extract_pdf(path: Path, max_chars: int) -> str:
    import fitz  # pymupdf

    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise AttachmentExtractionError(f"PDF konnte nicht geöffnet werden: {exc}") from exc

    if doc.is_encrypted:
        doc.close()
        raise EncryptedPdfError("PDF ist verschlüsselt und kann nicht gelesen werden.")

    parts: list[str] = []
    try:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text().strip()
            if page_text:
                parts.append(f"[Seite {page_num}]\n{page_text}")
    finally:
        doc.close()

    if not parts:
        return ""

    return _truncate_text("\n\n".join(parts), max_chars)


def _extract_docx(path: Path, max_chars: int) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _truncate_text("\n".join(parts), max_chars)


def _extract_xlsx(path: Path, max_chars: int, max_sheets: int = 5) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet_name in wb.sheetnames[:max_sheets]:
            sheet = wb[sheet_name]
            sheet_lines: list[str] = [f"[Arbeitsblatt {sheet_name}]"]
            for row in sheet.iter_rows(values_only=True):
                values = [
                    str(cell).strip()
                    for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if values:
                    sheet_lines.append(" | ".join(values))
            if len(sheet_lines) > 1:
                parts.append("\n".join(sheet_lines))
    finally:
        wb.close()
    return _truncate_text("\n\n".join(parts), max_chars)


class DocumentExtractor:
    """Extrahiert Text aus unterstützten Dokumentformaten."""

    def extract(self, path: Path, extension: str, max_chars: int) -> str:
        ext = extension.lower()
        if ext in (".txt", ".md", ".csv"):
            return _read_text_file(path, max_chars)
        if ext == ".pdf":
            return _extract_pdf(path, max_chars)
        if ext == ".docx":
            return _extract_docx(path, max_chars)
        if ext == ".xlsx":
            return _extract_xlsx(path, max_chars)
        raise AttachmentExtractionError(f"Nicht unterstützter Dokumenttyp: {ext}")
